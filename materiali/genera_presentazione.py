#!/usr/bin/env python3
# Genera il documento Word di presentazione del progetto Eventi Ferrara.
# Tema coerente col sito (navy/oro) + screenshot delle sezioni.
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

CENTER = WD_ALIGN_PARAGRAPH.CENTER
NAVY   = RGBColor(0x12, 0x2c, 0x4a)   # titoli di sezione
ORO    = RGBColor(0xa8, 0x82, 0x3a)   # accento / titolo copertina (leggibile su bianco)
GRIGIO = RGBColor(0x6e, 0x7c, 0x8a)
SHOTS  = os.path.join(os.path.dirname(__file__), "screenshots")

doc = Document()

# stile base: font senza grazie + spaziatura paragrafo 0
normal = doc.styles["Normal"]
normal.font.name = "Avenir Next Condensed"
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)

def _nospace(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p

def titolo(testo, size, color=None, bold=True, align=None, after=0, before=0):
    p = doc.add_paragraph(); _nospace(p)
    if align: p.alignment = align
    r = p.add_run(testo); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def h1(t):
    p = titolo(t, 16, NAVY, before=6)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    return p

def para(t, **k):
    return _nospace(doc.add_paragraph(t))

def punto(t, sub=False):
    p = doc.add_paragraph(style="List Bullet" if not sub else "List Bullet 2")
    p.add_run(t); return _nospace(p)

def kv(chiave, valore):
    p = _nospace(doc.add_paragraph())
    r = p.add_run(chiave + ": "); r.bold = True
    p.add_run(valore)

def immagine(nome, larghezza=6.2, didascalia=None):
    path = os.path.join(SHOTS, nome)
    if not os.path.exists(path): return
    doc.add_picture(path, width=Inches(larghezza))
    doc.paragraphs[-1].alignment = CENTER
    doc.paragraphs[-1].paragraph_format.space_before = Pt(6)
    if didascalia:
        c = _nospace(doc.add_paragraph()); c.alignment = CENTER
        r = c.add_run(didascalia); r.italic = True
        r.font.size = Pt(9); r.font.color.rgb = GRIGIO

# ===== Copertina =====
titolo("EVENTI FERRARA", 30, ORO, align=CENTER, after=2, before=30)
titolo("Calendario condiviso degli eventi della città", 14, GRIGIO, bold=False,
       align=CENTER, after=2)
titolo("Uno strumento al servizio di albergatori e operatori turistici", 12, GRIGIO, bold=False,
       align=CENTER, after=14)
p = doc.add_paragraph(); p.alignment = CENTER
r = p.add_run("https://eventiferrara.github.io/"); r.bold = True; r.font.color.rgb = ORO; r.font.size = Pt(13)
p2 = doc.add_paragraph(); p2.alignment = CENTER
p2.add_run("Canale Telegram: @calendarioeventiferrara").font.color.rgb = GRIGIO

immagine("home.png", larghezza=6.3, didascalia="La home della web app")

# ===== 1. Cos'è =====
h1("1. Che cos'è")
para("Eventi Ferrara è una web app pubblica e gratuita che raccoglie in un unico calendario "
     "tutti gli eventi che si svolgono in città — concerti, mostre, fiere, congressi, "
     "manifestazioni sportive — insieme a una previsione delle presenze turistiche attese.")
para("Lo scopo è dare a chi lavora nel turismo (alberghi, B&B, ristoranti, servizi) uno "
     "strumento semplice per sapere in anticipo quando la città sarà più affollata, così da "
     "organizzarsi al meglio su prezzi e personale.")
para("Studiando l'andamento delle presenze negli anni, l'app aiuta inoltre a riconoscere i "
     "periodi e i giorni di maggiore domanda e a predisporre un piano prezzi che la anticipi, "
     "invece di rincorrerla.")

# ===== 2. L'idea / la logica =====
h1("2. La logica del progetto")
para("Il progetto si fonda su tre idee:")
punto("Calendario collaborativo: nessuno conosce tutti gli eventi della città. Per questo "
      "ogni operatore può inserire gli eventi che conosce, arricchendo un calendario comune "
      "da cui tutti traggono beneficio.")
punto("Dato condiviso e sempre disponibile: le informazioni non restano sul singolo computer "
      "ma vivono su un database online, accessibile a tutti e salvato in modo sicuro.")
punto("Dall'informazione alla decisione: incrociando gli eventi con le presenze turistiche "
      "giornaliere ufficiali, l'app aiuta a capire l'impatto reale di un evento e a fare "
      "previsioni confrontando periodi e anni diversi.")

# ===== 3. Come funziona (operatore) =====
h1("3. Come funziona — per gli operatori")
para("La web app si apre nel browser (telefono o computer), senza installare nulla, ed è "
     "organizzata in quattro sezioni.")

titolo("Calendario eventi", 12, NAVY, after=2)
punto("Una tabella scorrevole mostra, giorno per giorno, gli eventi in programma, con un "
      "pallino colorato che indica il livello di presenze atteso (giallo = basso, verde = "
      "medio, rosso = alto).")
punto("Le pill in alto filtrano gli eventi per tipologia; le festività compaiono con il filtro “Tutti”.")
punto("Si può anche cercare una data specifica, anche a distanza di mesi o anni.")
immagine("calendario.png", larghezza=6.0, didascalia="Calendario eventi con previsione delle presenze")

titolo("Inserisci evento", 12, NAVY, after=2)
punto("Si indicano date di inizio e fine, nome dell'evento, tipologia (culturale/musicale, "
      "congressuale/fieristico, sportivo) e la previsione di presenze: bassa, media o alta.")
punto("Un controllo automatico segnala se un evento simile è già stato inserito, per evitare "
      "doppioni.")
immagine("inserisci.png", larghezza=6.0, didascalia="Form di inserimento di un nuovo evento")

titolo("Analisi presenze / evento", 12, NAVY, after=2)
punto("Confronto giorno per giorno: si sceglie un periodo e lo si confronta con lo stesso "
      "periodo dell'anno precedente, con un altro periodo a scelta oppure con un evento "
      "specifico, vedendo la differenza di presenze per ogni giorno.")
punto("Istogramma: un grafico a barre confronta le presenze di due periodi, con le stesse "
      "opzioni di confronto.")
punto("Pattern settimanale: scelto un mese, l'app mostra una linea per ogni anno (allineata "
      "per giorno della settimana, lunedì con lunedì) e una linea più spessa che ne indica la "
      "mediana, cioè il pattern “tipico”. Si vede così a colpo d'occhio quali giorni "
      "sono storicamente i più pieni e quanto è regolare l'andamento, base ideale per "
      "predisporre un piano prezzi che anticipa la domanda.")
immagine("analisi.png", larghezza=6.0, didascalia="Strumenti di analisi e confronto delle presenze")

titolo("Amministratore", 12, NAVY, after=2)
punto("Sezione riservata, protetta da accesso personale, da cui si caricano i file ufficiali "
      "delle presenze turistiche giornaliere (un file per anno, che viene aggiornato quando la "
      "Regione pubblica nuovi dati) e, volendo, gli eventi storici degli anni passati.")
punto("Da qui è anche possibile correggere o eliminare qualsiasi evento.")

# ===== 4. Telegram =====
h1("4. Avvisi automatici su Telegram")
para("Il sistema controlla ogni 30 minuti se sono stati inseriti nuovi eventi e, in tal caso, "
     "li pubblica automaticamente sul canale Telegram del progetto. In questo modo un nuovo "
     "evento viene segnalato sul canale entro circa mezz'ora dall'inserimento. Gli operatori "
     "che si iscrivono al canale restano così aggiornati senza dover controllare di continuo "
     "la web app.")

# ===== 5. Architettura =====
h1("5. Come è fatto (in breve)")
para("La soluzione usa servizi affidabili e gratuiti:")
punto("Web app pubblicata su GitHub Pages (l'indirizzo del sito).")
punto("Database online su Google Firebase: i dati sono salvati lato server a ogni modifica, "
      "non sul dispositivo dell'utente.")
punto("Un processo automatico salva periodicamente una copia di sicurezza dei dati "
      "(ripristinabile in qualsiasi momento) e invia gli avvisi sul canale Telegram.")
para("Sicurezza: chiunque può consultare il calendario e aggiungere eventi, ma solo "
     "l'amministratore può caricare i dati delle presenze e modificare o eliminare gli eventi.")

# ===== 6. Come si accede =====
h1("6. Come accedere")
kv("Sito web", "https://eventiferrara.github.io/")
kv("Canale Telegram", "@calendarioeventiferrara")
para("Per iniziare basta aprire il sito, consultare il calendario e inserire gli eventi che "
     "si conoscono. Più operatori partecipano, più il calendario diventa utile per tutti.")

doc.save("PRESENTAZIONE_Eventi_Ferrara.docx")
print("Creato PRESENTAZIONE_Eventi_Ferrara.docx")
